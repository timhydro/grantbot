from __future__ import annotations

import hashlib
import json
import os
import re
import tempfile
from dataclasses import asdict, dataclass
from datetime import date, datetime, timezone
from pathlib import Path
from typing import Any, Mapping, Sequence
from urllib.parse import urlparse

from grantbot.applications.review_workspace import application_root
from grantbot.knowledge.writer_facts import load_canonical_writer_facts
from grantbot.writing.quality_gate import evaluate_draft


TRUSTED_STATES = {"VERIFIED", "APPROVED"}
PLANNING_STATES = {"UNVERIFIED", "DRAFT"}
SAFE_ID_RE = re.compile(r"[^A-Za-z0-9._-]+")
SUPPORTED_DRAFT_MODES = {"template", "local_ai"}


@dataclass(frozen=True, slots=True)
class ResolvedField:
    key: str
    label: str
    value: Any
    status: str
    source: str
    fact_key: str | None
    required: bool
    hard_gate: bool
    instructions: str

    def to_dict(self) -> dict[str, Any]:
        return asdict(self)


@dataclass(frozen=True, slots=True)
class EligibilityResult:
    check_id: str
    label: str
    requirement: str
    status: str
    evidence: list[dict[str, Any]]
    source_url: str
    action: str

    def to_dict(self) -> dict[str, Any]:
        return asdict(self)


@dataclass(frozen=True, slots=True)
class SectionResult:
    section_id: str
    label: str
    prompt: str
    prompt_source: str
    section_type: str
    max_words: int | None
    status: str
    draft: str | None
    word_count: int
    facts_used: list[dict[str, Any]]
    missing_fact_keys: list[str]
    quality: dict[str, Any] | None
    drafting_method: str

    def to_dict(self) -> dict[str, Any]:
        return asdict(self)


def _safe_id(value: str) -> str:
    cleaned = SAFE_ID_RE.sub("_", str(value).strip()).strip("._-")
    if not cleaned:
        raise ValueError("opportunity_id cannot be empty")
    return cleaned[:120]


def _atomic_write(path: Path, text: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    fd, temporary = tempfile.mkstemp(prefix=path.name, suffix=".tmp", dir=str(path.parent))
    try:
        with os.fdopen(fd, "w", encoding="utf-8") as handle:
            handle.write(text)
            handle.flush()
            os.fsync(handle.fileno())
        os.replace(temporary, path)
    finally:
        if os.path.exists(temporary):
            os.unlink(temporary)


def _json_text(value: Any) -> str:
    return json.dumps(value, ensure_ascii=False, sort_keys=True, default=str)


def _value_text(value: Any) -> str:
    if value in (None, "", [], {}):
        return ""
    if isinstance(value, str):
        return " ".join(value.split())
    if isinstance(value, Sequence) and not isinstance(value, (str, bytes, bytearray)):
        values = [" ".join(str(item).split()) for item in value if str(item).strip()]
        if not values:
            return ""
        if len(values) == 1:
            return values[0]
        if len(values) == 2:
            return f"{values[0]} and {values[1]}"
        return ", ".join(values[:-1]) + f", and {values[-1]}"
    return " ".join(str(value).split())


def _validate_url(value: str, label: str) -> str:
    url = str(value or "").strip()
    parsed = urlparse(url)
    if parsed.scheme != "https" or not parsed.netloc:
        raise ValueError(f"{label} must be an absolute https URL")
    return url


def _validate_opportunity(raw: Mapping[str, Any]) -> dict[str, Any]:
    opportunity = dict(raw)
    for key in ("opportunity_id", "opportunity_name", "funder", "funding_type", "source_url"):
        value = str(opportunity.get(key) or "").strip()
        if not value:
            raise ValueError(f"{key} is required")
        opportunity[key] = value
    opportunity["opportunity_id"] = _safe_id(opportunity["opportunity_id"])
    opportunity["source_url"] = _validate_url(opportunity["source_url"], "source_url")

    deadline = str(opportunity.get("deadline") or "").strip()
    if deadline:
        try:
            datetime.strptime(deadline, "%Y-%m-%d")
        except ValueError as exc:
            raise ValueError("deadline must use YYYY-MM-DD") from exc
        opportunity["deadline"] = deadline
    else:
        opportunity["deadline"] = None

    for key in (
        "official_sources",
        "external_evidence",
        "priorities",
        "eligibility_checks",
        "required_portal_fields",
        "application_questions",
        "attachment_checklist",
        "submission_steps",
    ):
        value = opportunity.get(key, [])
        if not isinstance(value, list):
            raise ValueError(f"{key} must be a list")
        opportunity[key] = value

    if not opportunity["application_questions"]:
        raise ValueError("application_questions cannot be empty")

    for index, source in enumerate(opportunity["official_sources"]):
        if not isinstance(source, Mapping):
            raise ValueError(f"official_sources[{index}] must be an object")
        _validate_url(str(source.get("url") or ""), f"official_sources[{index}].url")

    for index, evidence in enumerate(opportunity["external_evidence"]):
        if not isinstance(evidence, Mapping):
            raise ValueError(f"external_evidence[{index}] must be an object")
        if not str(evidence.get("id") or "").strip():
            raise ValueError(f"external_evidence[{index}].id is required")
        _validate_url(
            str(evidence.get("source_url") or ""),
            f"external_evidence[{index}].source_url",
        )

    request = opportunity.get("request") or {}
    if not isinstance(request, Mapping):
        raise ValueError("request must be an object")
    opportunity["request"] = dict(request)
    return opportunity


def load_opportunity_file(path: Path) -> dict[str, Any]:
    source = Path(path).expanduser().resolve()
    data = json.loads(source.read_text(encoding="utf-8"))
    if not isinstance(data, Mapping):
        raise ValueError("opportunity input must contain a JSON object")
    return _validate_opportunity(data)


def builtin_opportunity_path(name: str) -> Path:
    safe = _safe_id(name)
    path = Path(__file__).resolve().parents[1] / "opportunities" / f"{safe}.json"
    if not path.is_file():
        raise FileNotFoundError(f"Built-in opportunity not found: {safe}")
    return path


def _normalize_fact(raw: Mapping[str, Any]) -> dict[str, Any]:
    state = str(raw.get("verification_state") or raw.get("status") or "MISSING").upper()
    writer_state = "DRAFT" if state in PLANNING_STATES else state
    return {
        "id": str(raw.get("fact_id") or raw.get("id") or ""),
        "category": str(raw.get("category") or "general").lower(),
        "key": str(raw.get("fact_key") or raw.get("key") or "").lower(),
        "value": raw.get("value", raw.get("answer")),
        "status": writer_state,
        "verification_state": state,
        "source": str(raw.get("source") or "unspecified"),
        "source_type": str(raw.get("source_type") or ""),
        "confidence": float(raw.get("confidence") or 0.0),
        "notes": str(raw.get("notes") or ""),
    }


def canonical_application_facts() -> list[dict[str, Any]]:
    return load_canonical_writer_facts(actor="production-application-packet-v30.1")


def _fact_index(facts: Sequence[Mapping[str, Any]]) -> dict[str, dict[str, Any]]:
    return {str(item.get("key") or "").lower(): dict(item) for item in facts if item.get("key")}


def _best_fact(
    index: Mapping[str, Mapping[str, Any]],
    keys: Sequence[str],
) -> dict[str, Any] | None:
    candidates = [
        dict(index[key.strip().lower()])
        for key in keys
        if key.strip().lower() in index
        and index[key.strip().lower()].get("value") not in (None, "", [], {})
    ]
    candidates.sort(
        key=lambda item: (
            0 if str(item.get("verification_state")) == "VERIFIED" else
            1 if str(item.get("verification_state")) == "APPROVED" else
            2 if str(item.get("verification_state")) in PLANNING_STATES else 3,
            -float(item.get("confidence") or 0),
        )
    )
    return candidates[0] if candidates else None


def _resolve_portal_fields(
    specs: Sequence[Mapping[str, Any]],
    facts: Sequence[Mapping[str, Any]],
) -> list[ResolvedField]:
    index = _fact_index(facts)
    output: list[ResolvedField] = []
    for position, spec in enumerate(specs):
        key = str(spec.get("key") or "").strip().lower()
        label = str(spec.get("label") or key).strip()
        if not key or not label:
            raise ValueError(f"required_portal_fields[{position}] requires key and label")
        aliases = [str(item) for item in spec.get("fact_keys", [key]) if str(item).strip()]
        fact = _best_fact(index, aliases)
        required = bool(spec.get("required", True))
        hard_gate = bool(spec.get("hard_gate", required))
        instructions = str(spec.get("instructions") or "").strip()
        if fact is None:
            status = "MISSING" if required else "OPTIONAL_NOT_SUPPLIED"
            value = None
            source = "Canonical knowledge: no usable fact"
            fact_key = None
        else:
            state = str(fact.get("verification_state") or "MISSING")
            value = fact.get("value")
            source = str(fact.get("source") or "unspecified")
            fact_key = str(fact.get("key") or "") or None
            if state in TRUSTED_STATES:
                status = "READY"
            else:
                status = "CONFIRM" if required else "OPTIONAL_CONFIRM"
        output.append(ResolvedField(
            key=key,
            label=label,
            value=value,
            status=status,
            source=source,
            fact_key=fact_key,
            required=required,
            hard_gate=hard_gate,
            instructions=instructions,
        ))
    return output


def _eligibility_results(
    specs: Sequence[Mapping[str, Any]],
    facts: Sequence[Mapping[str, Any]],
) -> list[EligibilityResult]:
    index = _fact_index(facts)
    output: list[EligibilityResult] = []
    for position, spec in enumerate(specs):
        check_id = _safe_id(str(spec.get("id") or f"check-{position + 1}"))
        label = str(spec.get("label") or check_id).strip()
        requirement = str(spec.get("requirement") or "").strip()
        source_url = _validate_url(str(spec.get("source_url") or ""), f"eligibility_checks[{position}].source_url")
        keys = [str(item).strip().lower() for item in spec.get("evidence_fact_keys", []) if str(item).strip()]
        mode = str(spec.get("mode") or "any").strip().lower()
        if mode not in {"any", "all"}:
            raise ValueError("eligibility check mode must be any or all")
        evidence: list[dict[str, Any]] = []
        for key in keys:
            item = index.get(key)
            if item and item.get("value") not in (None, "", [], {}):
                evidence.append({
                    "fact_key": key,
                    "value": item.get("value"),
                    "verification_state": item.get("verification_state"),
                    "source": item.get("source"),
                })
        trusted_evidence = [
            item
            for item in evidence
            if str(item.get("verification_state")) in TRUSTED_STATES
        ]
        trusted_keys = {
            str(item["fact_key"])
            for item in trusted_evidence
        }
        planning_present = any(
            str(item.get("verification_state")) in PLANNING_STATES
            for item in evidence
        )
        evidence_satisfied = (
            bool(trusted_keys)
            if mode == "any"
            else bool(keys) and all(key in trusted_keys for key in keys)
        )
        semantic_text = _json_text(
            [item.get("value") for item in trusted_evidence]
        ).lower()
        required_any = [
            str(item).strip().lower()
            for item in spec.get("required_substrings_any", [])
            if str(item).strip()
        ]
        required_all = [
            str(item).strip().lower()
            for item in spec.get("required_substrings_all", [])
            if str(item).strip()
        ]
        semantic_satisfied = (
            (not required_any or any(item in semantic_text for item in required_any))
            and all(item in semantic_text for item in required_all)
        )
        if evidence_satisfied and semantic_satisfied:
            status = "PASS"
            action = "Retain source evidence in the final review file."
        elif evidence_satisfied:
            status = "FAIL" if bool(spec.get("hard_fail_on_mismatch")) else "NEEDS_CONFIRMATION"
            action = str(
                spec.get("action")
                or "The recorded fact does not establish this requirement; verify eligibility before proceeding."
            )
        elif planning_present:
            status = "NEEDS_CONFIRMATION"
            action = str(spec.get("action") or "Confirm the planning fact before final eligibility approval.")
        else:
            status = "NEEDS_INFORMATION"
            action = str(spec.get("action") or "Supply authoritative applicant evidence before proceeding.")
        output.append(EligibilityResult(
            check_id=check_id,
            label=label,
            requirement=requirement,
            status=status,
            evidence=evidence,
            source_url=source_url,
            action=action,
        ))
    return output


def _deadline_review(deadline: str | None) -> dict[str, Any]:
    if not deadline:
        return {
            "deadline": None,
            "status": "NOT_STATED",
            "days_remaining": None,
            "review_date": date.today().isoformat(),
        }
    deadline_date = datetime.strptime(deadline, "%Y-%m-%d").date()
    today = date.today()
    days_remaining = (deadline_date - today).days
    if days_remaining < 0:
        status = "EXPIRED"
    elif days_remaining == 0:
        status = "DUE_TODAY"
    else:
        status = "OPEN"
    return {
        "deadline": deadline,
        "status": status,
        "days_remaining": days_remaining,
        "review_date": today.isoformat(),
    }


def _fact_value(index: Mapping[str, Mapping[str, Any]], *keys: str) -> str:
    fact = _best_fact(index, list(keys))
    return _value_text(fact.get("value")) if fact else ""


def _external_evidence_facts(opportunity: Mapping[str, Any]) -> list[dict[str, Any]]:
    output: list[dict[str, Any]] = []
    for raw in opportunity.get("external_evidence", []):
        evidence_id = _safe_id(str(raw.get("id") or "evidence"))
        output.append(
            {
                "id": f"external-{evidence_id}",
                "category": "external_evidence",
                "key": f"external_{evidence_id.lower()}",
                "value": raw.get("source_values") or raw.get("application_language") or raw,
                "status": "VERIFIED",
                "verification_state": "VERIFIED",
                "source": str(raw.get("source_url") or ""),
                "source_type": "official_external_report",
                "confidence": 1.0,
                "notes": str(raw.get("caveat") or ""),
                "external_evidence_id": evidence_id,
            }
        )
    return output


def _external_evidence(
    opportunity: Mapping[str, Any],
    evidence_id: str,
) -> Mapping[str, Any] | None:
    expected = _safe_id(evidence_id)
    for item in opportunity.get("external_evidence", []):
        if _safe_id(str(item.get("id") or "")) == expected:
            return item
    return None


def _prospective(text: str) -> str:
    value = text.strip()
    if not value:
        return ""
    return value[0].lower() + value[1:] if len(value) > 1 else value.lower()


def _display_status(value: str) -> str:
    labels = {
        "READY": "Ready",
        "MISSING": "Missing",
        "CONFIRM": "Confirm",
        "OPTIONAL_CONFIRM": "Optional — confirm",
        "OPTIONAL_NOT_SUPPLIED": "Optional",
        "READY_FOR_HUMAN_REVIEW": "Ready for human review",
        "NEEDS_INFORMATION": "Needs information",
        "REVISION_REQUIRED": "Revision required",
        "TO_CONFIRM_IN_PORTAL": "Confirm in live portal",
        "RECOMMENDED_INTERNAL_FILE": "Prepare if requested",
        "HUMAN_APPROVAL_REQUIRED": "Human approval required",
        "RECOMMENDED_IF_PROJECT_REQUEST": "Prepare for project request",
        "RECOMMENDED_IF_AVAILABLE": "Include if confirmed",
        "AVAILABLE": "Available",
        "NEEDS_USER": "Needs information",
        "READY_FOR_REVIEW": "Ready for review",
        "HOLD": "Hold",
    }
    normalized = str(value or "").strip().upper()
    return labels.get(normalized, normalized.replace("_", " ").title())


def _display_key(value: str) -> str:
    return str(value or "").replace("_", " ").strip()


def _human_date(value: str) -> str:
    try:
        return datetime.strptime(value, "%Y-%m-%d").strftime("%B %d, %Y").replace(" 0", " ")
    except ValueError:
        return value


def _trim_words(text: str, max_words: int | None) -> str:
    cleaned = " ".join(text.split())
    if max_words is None:
        return cleaned
    words = cleaned.split()
    if len(words) <= max_words:
        return cleaned
    return " ".join(words[:max_words]).rstrip(" ,;:") + "."


def _deterministic_draft(
    *,
    section_type: str,
    opportunity: Mapping[str, Any],
    facts: Sequence[Mapping[str, Any]],
    max_words: int | None,
) -> str:
    index = _fact_index(facts)
    legal_name = _fact_value(index, "legal_name")
    organization = "Broken Growth Ministries"
    mission = _fact_value(index, "mission_statement", "vision_statement")
    mission = mission.replace("BrokenGrowthMinistries", organization, 1)
    geography = _fact_value(index, "primary_geographic_focus", "service_area")
    populations = _fact_value(index, "primary_populations")
    faith = _fact_value(index, "faith_identity")
    program = _fact_value(index, "program_model")
    housing = _fact_value(index, "planned_tiny_homes", "tiny_home_model")
    shelter_model = _fact_value(index, "pallet_shelter_model")
    project_site = _fact_value(index, "project_site_address")
    property_status = _fact_value(index, "property_status")
    capital_plan = _fact_value(index, "capital_planning_budget")
    projected = _fact_value(index, "projected_people_served", "annual_people_served")
    capacity = _fact_value(index, "housing_capacity")
    eligibility = _fact_value(index, "eligibility_requirements")
    sobriety = _fact_value(index, "sobriety_requirement")
    stabilization = _fact_value(index, "stabilization_period")
    leadership = _fact_value(index, "executive_director", "founder_name")
    board_size = _fact_value(index, "board_size")
    incorporated = _fact_value(index, "date_incorporated")
    entity_status = _fact_value(index, "state_entity_status")
    funding_scope = _fact_value(index, "funding_source_scope")
    resident_contribution = _fact_value(index, "resident_contribution_model", "resident_fees")
    services_status = _fact_value(index, "current_service_status")
    funder = str(opportunity.get("funder") or "the funder")
    program_name = str((opportunity.get("request") or {}).get("project_name") or "the first-phase housing and shelter program")
    pit = _external_evidence(opportunity, "fl511-pit-2026")
    pit_summary = str((pit or {}).get("application_language") or "").strip()
    pit_short = str((pit or {}).get("executive_language") or "").strip()

    if section_type == "executive_summary":
        text = (
            f"{organization} seeks a partnership with {funder} to advance {program_name} in {geography}. "
            f"{mission} The proposed work aligns housing with practical stabilization through {_prospective(program).rstrip('.')}. "
            f"{pit_short} "
            f"Support would help move the ministry from early community service and program planning toward a safe, accountable housing response for {populations}. "
            f"The request is intentionally focused on housing and shelter readiness, while employment, life skills, mentorship, and resource navigation create a pathway from immediate safety toward durable independence."
        )
    elif section_type == "statement_of_need":
        text = (
            f"Broken Growth is designed for {populations} in {geography}. {pit_summary} "
            "For people leaving incarceration or living without stable shelter, housing insecurity can disrupt employment, recovery, family stability, and successful community reentry at the same time. "
            f"The ministry's response begins with safe housing and pairs it with {_prospective(program).rstrip('.')}. This integrated design addresses the practical conditions that make stability difficult to sustain, rather than treating shelter as an isolated service. "
            f"The program will use verified local data and documented participant records as they become available; this draft does not claim outcomes or community statistics that Broken Growth has not yet substantiated."
        )
    elif section_type == "program_design":
        text = (
            f"The planned program will serve {populations} through {_prospective(program).rstrip('.')}. "
            f"The physical model is {_prospective(shelter_model or housing).rstrip('.')}. "
            f"Admission and stabilization will follow the approved program policy: {eligibility} "
            "After move-in, participants will receive a 30-day stabilization period before work, training, or active job-search requirements begin. "
            f"This sequence protects the first weeks of housing as a period for stabilization and individualized planning before employment or training expectations begin. Staff and partners will then coordinate housing stability, workforce preparation, life skills, mentorship, transportation assistance where feasible, and resource navigation. "
            f"Every application and participant record will remain subject to human review, safety procedures, and documented program rules."
        )
    elif section_type == "goals_objectives":
        text = (
            f"Broken Growth's first-phase goal is a safe, structured housing pathway for {populations}. Planned capacity is {housing}, with {projected}. "
            "Measurable program targets will include shelters made ready, participants enrolled, individualized plans completed, referrals and training connections made, and housing and employment status at scheduled follow-up points. "
            "A second objective is to connect housing stability to employment, training, life skills, mentorship, and reentry support. A third objective is consistent documentation of participation and progress so actual outputs and outcomes are never confused with projections. "
            "Implementation dependencies include an approved request budget, a verified property pathway, agency and engineering decisions, staffing and safety plans, current quotes, and a board-approved implementation schedule. Final baselines, numeric targets, and dates will be approved only after those dependencies are resolved."
        )
        if capacity and projected:
            projected_target = re.sub(
                r"^first-year planning target of\s+",
                "",
                _prospective(projected).rstrip("."),
                flags=re.IGNORECASE,
            )
            text = text.replace(
                f"Planned capacity is {housing}, with {projected}.",
                "Planned capacity is "
                f"{_prospective(capacity).rstrip('.')}; the first-year service plan is "
                f"to reach a planning target of {projected_target}.",
            )
    elif section_type == "outcomes_evaluation":
        text = (
            "Broken Growth will separate service outputs from participant outcomes. Outputs will include enrollments, housing placements, service-plan completion, referrals, training participation, employment connections, and follow-up contacts. Outcomes will focus on housing retention, employment retention, income stability, sobriety and program compliance where applicable, successful community reentry, and movement toward independent living. "
            "The ministry will use intake records, individualized plans, attendance and referral logs, housing and employment status checks, and scheduled follow-up reviews. Baselines and targets will be established before launch and approved by program leadership; no projected result will be reported as an achieved outcome."
        )
    elif section_type == "organizational_capacity":
        text = (
            f"{organization} ({legal_name}) is an {entity_status.lower() if entity_status else 'active'} Florida nonprofit corporation incorporated on {_human_date(incorporated)}. It is led by Founder and Executive Director {leadership} with a board of {board_size}. "
            f"The organization is a {faith} rooted in a mission of second chances, safe housing, meaningful work, supportive community, independence, and personal growth. {services_status} "
            "Because the organization is still building its housing platform, this application distinguishes current activity from planned capacity and routes legal, financial, property, safety, and submission decisions to authorized human reviewers."
        )
    elif section_type == "project_readiness":
        site = project_site.replace("Proposed project site:", "").strip()
        text = (
            f"Broken Growth is preparing {program_name} for the proposed project site at {site}. The site remains in predevelopment planning; site control, zoning, use classification, utilities, buildable capacity, and development approvals require human and agency verification. "
            f"The working concept is {_prospective(shelter_model).rstrip('.')}. {capital_plan} "
            "Before committing grant funds or representing the project as construction-ready, the organization will confirm site control, zoning and use classification, legal unit capacity, utilities, accessibility, fire and life safety, stormwater, engineering, vendor scope, current quotes, insurance, the implementation schedule, and board-approved financing. "
            "Planning assumptions will remain labeled as assumptions, and each will be replaced with written agency, professional, vendor, or governing-board evidence as it becomes available. This staged approach allows Publix support to be matched to a defined, achievable phase without overstating readiness."
        )
    elif section_type == "sustainability":
        text = (
            "Broken Growth plans a diversified sustainability strategy using public grants, foundation and corporate philanthropy, faith and community support, program and operating grants, capital funding, matching support, and documented in-kind contributions. "
            f"The service model also anticipates {_prospective(resident_contribution).rstrip('.')}. "
            "Grant funds would be used to create credible early capacity and documentation rather than assuming permanent dependence on one funder. The ministry will pair disciplined budgeting and outcome tracking with continued public, private, faith-based, corporate, and community fundraising. Final sustainability claims will be updated after the operating budget, confirmed partners, and revenue assumptions receive board approval."
        )
    elif section_type == "budget_narrative":
        text = (
            f"The proposed budget for {program_name} will be tied directly to verified housing and shelter implementation costs. Allowable uses will be mapped to the final scope, such as site and unit readiness, safety-related improvements, furnishings or equipment, program delivery, participant stabilization supports, and other costs specifically permitted by {funder}. "
            "Broken Growth will not enter a requested amount or line-item total until quotations, funding restrictions, organizational financial records, and board authorization are complete. The final budget narrative will reconcile exactly to the approved budget and will not include unsupported estimates."
        )
    elif section_type == "partnerships":
        text = (
            "Broken Growth's model depends on coordinated relationships with housing, workforce, reentry, behavioral-health, faith, transportation, employer, and community-service partners. The application will describe only partners and roles that are documented and authorized. Prospective relationships will be labeled as prospective, and letters of commitment will be requested where the funder or implementation plan requires them."
        )
    else:
        text = (
            f"{organization} serves {populations} in {geography}. {mission} The proposed request to {funder} supports {program_name} through {program} "
            "The final response will use only verified or founder-approved facts and will remain subject to human review before portal entry or submission."
        )
    return _trim_words(text, max_words)


def _section_facts(
    spec: Mapping[str, Any],
    facts: Sequence[Mapping[str, Any]],
) -> tuple[list[dict[str, Any]], list[str]]:
    index = _fact_index(facts)
    required = [str(item).strip().lower() for item in spec.get("required_fact_keys", []) if str(item).strip()]
    supporting = [str(item).strip().lower() for item in spec.get("supporting_fact_keys", []) if str(item).strip()]
    selected: list[dict[str, Any]] = []
    missing: list[str] = []
    for key in required:
        fact = _best_fact(index, [key])
        if fact is None or fact.get("verification_state") == "MISSING":
            missing.append(key)
        else:
            selected.append(fact)
    for key in supporting:
        fact = _best_fact(index, [key])
        if fact is not None and fact not in selected:
            selected.append(fact)
    if not selected:
        selected = [
            dict(item)
            for item in facts
            if str(item.get("verification_state")) in TRUSTED_STATES | PLANNING_STATES
            and item.get("value") not in (None, "", [], {})
        ]
    return selected, missing


def _build_dossier(opportunity: Mapping[str, Any], facts: Sequence[Mapping[str, Any]]) -> str:
    lines = [
        f"OPPORTUNITY: {opportunity['opportunity_name']}",
        f"FUNDER: {opportunity['funder']}",
        f"FUNDING TYPE: {opportunity['funding_type']}",
        f"DEADLINE: {opportunity.get('deadline') or 'not stated'}",
        f"OFFICIAL SOURCE: {opportunity['source_url']}",
        "",
        "FUNDER PRIORITIES:",
    ]
    lines.extend(f"- {item}" for item in opportunity.get("priorities", []))
    lines.extend(["", "VERIFIED EXTERNAL NEED EVIDENCE:"])
    for evidence in opportunity.get("external_evidence", []):
        lines.append(
            f"- {evidence.get('label')}: {evidence.get('source_values')} "
            f"(source: {evidence.get('source_url')}; caveat: {evidence.get('caveat')})"
        )
    lines.extend(["", "CANONICAL ORGANIZATION FACTS:"])
    for fact in facts:
        if fact.get("value") in (None, "", [], {}):
            continue
        lines.append(
            f"- [{fact.get('verification_state')}] {fact.get('category')}.{fact.get('key')}: "
            f"{_json_text(fact.get('value'))} (source: {fact.get('source')})"
        )
    lines.extend([
        "",
        "CONTROL RULES:",
        "- State verified and approved facts as facts.",
        "- Keep planning facts prospective.",
        "- Never invent missing legal, financial, outcome, partnership, property, or submission information.",
        "- Human review is required; no external submission is authorized.",
    ])
    return "\n".join(lines)


def _draft_sections(
    opportunity: Mapping[str, Any],
    facts: Sequence[Mapping[str, Any]],
    *,
    draft_mode: str,
    supplied_answers: Mapping[str, str] | None,
    run_adversarial_review: bool,
) -> list[SectionResult]:
    if draft_mode not in SUPPORTED_DRAFT_MODES:
        raise ValueError(f"draft_mode must be one of {', '.join(sorted(SUPPORTED_DRAFT_MODES))}")
    supplied = {str(key): str(value).strip() for key, value in (supplied_answers or {}).items() if str(value).strip()}
    dossier = _build_dossier(opportunity, facts)
    external_facts = {
        str(item.get("external_evidence_id")): item
        for item in _external_evidence_facts(opportunity)
    }
    output: list[SectionResult] = []
    for position, raw in enumerate(opportunity["application_questions"]):
        if not isinstance(raw, Mapping):
            raise ValueError(f"application_questions[{position}] must be an object")
        section_id = _safe_id(str(raw.get("id") or f"section-{position + 1}"))
        label = str(raw.get("label") or section_id).strip()
        prompt = str(raw.get("prompt") or "").strip()
        if not prompt:
            raise ValueError(f"application_questions[{position}].prompt is required")
        prompt_source = str(raw.get("prompt_source") or "PREPARATION_ONLY_NOT_PORTAL_VERBATIM").strip()
        section_type = str(raw.get("section_type") or "general").strip().lower()
        max_words_raw = raw.get("max_words")
        max_words = int(max_words_raw) if max_words_raw not in (None, "") else None
        if max_words is not None and max_words <= 0:
            raise ValueError("max_words must be positive")
        selected_facts, missing = _section_facts(raw, facts)
        for evidence_id in raw.get("external_evidence_ids", []):
            item = external_facts.get(_safe_id(str(evidence_id)))
            if item is not None and item not in selected_facts:
                selected_facts.append(item)

        method = draft_mode
        adaptive_status: str | None = None
        if section_id in supplied:
            draft = _trim_words(supplied[section_id], max_words)
            method = "operator_supplied"
        elif draft_mode == "local_ai":
            from grantbot.production.master_tool_v30 import run_adaptive_application_answer

            result = run_adaptive_application_answer(
                dossier=dossier,
                question=prompt,
                max_words=max_words,
                section_type=section_type,
                funder_archetype="corporate_foundation",
                run_adversarial_review=run_adversarial_review,
            )
            draft = str(result.get("final_answer") or "").strip()
            adaptive_status = str(result.get("adaptive_status") or "REMEDIATION_REQUIRED")
        else:
            draft = _deterministic_draft(
                section_type=section_type,
                opportunity=opportunity,
                facts=facts,
                max_words=max_words,
            )

        quality = None
        if draft:
            quality_result = evaluate_draft(
                question=prompt,
                draft=draft,
                facts=selected_facts,
                max_words=max_words,
            )
            quality = quality_result.to_dict()
        if missing:
            status = "NEEDS_INFORMATION"
        elif not draft:
            status = "NEEDS_INFORMATION"
        elif adaptive_status and adaptive_status != "READY_FOR_HUMAN_REVIEW":
            status = "REVISION_REQUIRED"
        elif quality and quality.get("passed"):
            status = "READY_FOR_HUMAN_REVIEW"
        else:
            status = "REVISION_REQUIRED"
        output.append(SectionResult(
            section_id=section_id,
            label=label,
            prompt=prompt,
            prompt_source=prompt_source,
            section_type=section_type,
            max_words=max_words,
            status=status,
            draft=draft or None,
            word_count=len(draft.split()) if draft else 0,
            facts_used=[
                {
                    "fact_key": item.get("key"),
                    "value": item.get("value"),
                    "verification_state": item.get("verification_state"),
                    "source": item.get("source"),
                }
                for item in selected_facts
            ],
            missing_fact_keys=missing,
            quality=quality,
            drafting_method=method,
        ))
    return output


def _attachment_results(
    specs: Sequence[Mapping[str, Any]],
    facts: Sequence[Mapping[str, Any]],
) -> list[dict[str, Any]]:
    index = _fact_index(facts)
    output: list[dict[str, Any]] = []
    for position, spec in enumerate(specs):
        item_id = _safe_id(str(spec.get("id") or f"attachment-{position + 1}"))
        label = str(spec.get("label") or item_id).strip()
        required = bool(spec.get("required", True))
        keys = [str(item).strip().lower() for item in spec.get("evidence_fact_keys", []) if str(item).strip()]
        fact = _best_fact(index, keys) if keys else None
        if fact and fact.get("verification_state") in TRUSTED_STATES:
            status = "AVAILABLE"
            evidence = {
                "fact_key": fact.get("key"),
                "value": fact.get("value"),
                "source": fact.get("source"),
            }
        elif fact:
            status = "CONFIRM"
            evidence = {
                "fact_key": fact.get("key"),
                "value": fact.get("value"),
                "source": fact.get("source"),
            }
        else:
            status = str(spec.get("source_status") or "TO_CONFIRM_IN_PORTAL") if not keys else "MISSING"
            evidence = None
        output.append({
            "id": item_id,
            "label": label,
            "required": required,
            "status": status,
            "evidence": evidence,
            "instructions": str(spec.get("instructions") or "").strip(),
            "source_status": str(spec.get("source_status") or "TO_CONFIRM_IN_PORTAL"),
        })
    return output


def _dedupe(items: Sequence[str]) -> list[str]:
    return list(dict.fromkeys(" ".join(str(item).split()) for item in items if str(item).strip()))


def _readiness_score(
    fields: Sequence[ResolvedField],
    eligibility: Sequence[EligibilityResult],
    sections: Sequence[SectionResult],
    attachments: Sequence[Mapping[str, Any]],
) -> int:
    required_fields = [item for item in fields if item.required]
    field_score = (
        sum(item.status == "READY" for item in required_fields) / len(required_fields)
        if required_fields else 1.0
    )
    eligibility_score = (
        sum(item.status == "PASS" for item in eligibility) / len(eligibility)
        if eligibility else 1.0
    )
    section_score = (
        sum(item.status == "READY_FOR_HUMAN_REVIEW" for item in sections) / len(sections)
        if sections else 0.0
    )
    required_attachments = [item for item in attachments if bool(item.get("required"))]
    attachment_score = (
        sum(item.get("status") == "AVAILABLE" for item in required_attachments) / len(required_attachments)
        if required_attachments else 1.0
    )
    return max(0, min(100, round(
        field_score * 40 + eligibility_score * 25 + section_score * 25 + attachment_score * 10
    )))


def _markdown_escape(value: Any) -> str:
    return _value_text(value).replace("|", "\\|").replace("\n", " ")


def _render_application_markdown(packet: Mapping[str, Any]) -> str:
    opportunity = packet["opportunity"]
    request = opportunity.get("request") or {}
    lines = [
        f"# {opportunity['opportunity_name']}",
        "",
        "## Human-review application draft",
        "",
        f"**Applicant:** {_markdown_escape(packet['organization_name'])}",
        f"**Funder:** {_markdown_escape(opportunity['funder'])}",
        f"**Funding type:** {_markdown_escape(opportunity['funding_type'])}",
        f"**Deadline:** {_markdown_escape(opportunity.get('deadline') or 'Not stated')}",
        f"**Requested amount:** {_markdown_escape(request.get('amount') or 'Not yet approved')}",
        f"**GrantBot readiness:** {packet['readiness_score']}%",
        f"**Packet status:** {packet['status']}",
        f"**Official source:** {opportunity['source_url']}",
        "",
        "> This is a GrantBot preparation and human-review packet. The published portal questions must be checked after login. Nothing in this file is authorized for automatic submission, certification, signature, payment, debt, or financial commitment.",
        "",
        "## Eligibility review",
        "",
        "| Requirement | Status | Evidence / action |",
        "|---|---:|---|",
    ]
    for item in packet["eligibility"]:
        evidence = "; ".join(
            f"{_display_key(entry['fact_key'])}: {_value_text(entry['value'])}"
            for entry in item["evidence"]
        ) or item["action"]
        lines.append(f"| {_markdown_escape(item['label'])} | {_display_status(item['status'])} | {_markdown_escape(evidence)} |")

    lines.extend(["", "## Verified external need evidence", "", "| Evidence | Source values | Scope caveat |", "|---|---|---|"])
    for item in opportunity["external_evidence"]:
        lines.append(
            f"| {_markdown_escape(item['label'])} | {_markdown_escape(item.get('source_values'))} | "
            f"{_markdown_escape(item.get('caveat'))} |"
        )

    lines.extend(["", "## Registration and portal fields", "", "| Field | Status | Prepared value | Entry instruction |", "|---|---:|---|---|"])
    for item in packet["portal_fields"]:
        value = _markdown_escape(item["value"]) or "Not supplied - action required"
        instructions = _markdown_escape(item["instructions"]) or "Verify against the live portal."
        lines.append(f"| {_markdown_escape(item['label'])} | {_display_status(item['status'])} | {value} | {instructions} |")

    lines.extend(["", "## Prepared narratives", ""])
    for section in packet["sections"]:
        lines.extend([
            f"### {section['label']}",
            "",
            f"**Status:** {_display_status(section['status'])}  ",
            f"**Preparation prompt:** {_markdown_escape(section['prompt'])}  ",
            f"**Prompt source:** {_markdown_escape(section['prompt_source'])}",
            "",
            section.get("draft") or "No narrative generated because required information is missing.",
            "",
        ])
        if section["missing_fact_keys"]:
            lines.append("Information still required: " + ", ".join(_display_key(item) for item in section["missing_fact_keys"]))
            lines.append("")

    lines.extend(["## Human-required gates", ""])
    lines.extend(f"- [ ] {item}" for item in packet["hard_gates"])
    if not packet["hard_gates"]:
        lines.append("- [x] No unresolved hard gates were recorded; authorized final review is still required.")

    lines.extend(["", "## Document preparation checklist", ""])
    for item in packet["attachments"]:
        lines.append(f"- [ ] {item['label']} - {_display_status(item['status'])}")

    lines.extend(["", "## Submission sequence", ""])
    lines.extend(f"{position}. {step}" for position, step in enumerate(packet["submission_steps"], start=1))

    lines.extend(["", "## Official sources", ""])
    for source in opportunity["official_sources"]:
        lines.append(f"- [{source.get('title') or source['url']}]({source['url']}) - {_value_text(source.get('supports'))}")
    lines.extend([
        "",
        "## Control statement",
        "",
        "GrantBot may research, score, draft, check, package, and stage this request. An authorized human must verify the live portal, resolve every hard gate, approve the final budget and narrative, make all certifications, and perform the actual submission.",
        "",
    ])
    return "\n".join(lines)


def _render_checklist_markdown(packet: Mapping[str, Any]) -> str:
    lines = [
        f"# Finalization Checklist - {packet['opportunity']['opportunity_name']}",
        "",
        f"Current readiness: {packet['readiness_score']}%",
        f"Current status: {packet['status']}",
        "",
        "## Resolve first",
        "",
    ]
    lines.extend(f"- [ ] {item}" for item in packet["hard_gates"])
    if not packet["hard_gates"]:
        lines.append("- [x] No recorded hard gates remain.")
    lines.extend([
        "",
        "## Final human review",
        "",
        "- [ ] Log in on a supported desktop browser and copy the exact live portal prompts into GrantBot.",
        "- [ ] Map each prepared narrative to the matching portal field; do not assume the preparation prompts are verbatim.",
        "- [ ] Confirm every organization, contact, legal, tax, financial, property, and program fact.",
        "- [ ] Reconcile every dollar and line item to the board-approved request budget.",
        "- [ ] Confirm all attachments, character limits, certifications, and the portal deadline/timezone.",
        "- [ ] Have an authorized officer approve the final application before submission.",
        "- [ ] Save the portal confirmation and submitted copy in GrantBot.",
        "",
    ])
    return "\n".join(lines)


def _set_table_geometry(table: Any, widths_dxa: Sequence[int], *, indent_dxa: int = 120) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    if sum(widths_dxa) != 9360:
        raise ValueError("table widths must total 9360 DXA")
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    for tag in ("w:tblW", "w:tblInd"):
        existing = tbl_pr.find(qn(tag))
        if existing is not None:
            tbl_pr.remove(existing)
    tbl_w = OxmlElement("w:tblW")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), "9360")
    tbl_pr.append(tbl_w)
    tbl_ind = OxmlElement("w:tblInd")
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_pr.append(tbl_ind)
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths_dxa, strict=True):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(width))


def _shade_cell(cell: Any, fill: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def _set_cell_margins(cell: Any, *, top: int = 80, bottom: int = 80, start: int = 120, end: int = 120) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.find(qn("w:tcMar"))
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for edge, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        node = margins.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_run_font(run: Any, *, name: str = "Calibri", size: float | None = None, color: str | None = None, bold: bool | None = None, italic: bool | None = None) -> None:
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor

    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def _configure_document_styles(document: Any) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt, RGBColor

    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    tokens = {
        "Heading 1": (16, "2E74B5", 16, 8),
        "Heading 2": (13, "2E74B5", 12, 6),
        "Heading 3": (12, "1F4D78", 8, 4),
    }
    for name, (size, color, before, after) in tokens.items():
        style = document.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = document.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.194)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.208
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT


def _add_table(document: Any, headers: Sequence[str], rows: Sequence[Sequence[str]], widths: Sequence[int]) -> Any:
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    header = table.rows[0]
    header_properties = header._tr.get_or_add_trPr()
    repeat_header = OxmlElement("w:tblHeader")
    repeat_header.set(qn("w:val"), "true")
    header_properties.append(repeat_header)
    for cell, label in zip(header.cells, headers, strict=True):
        _shade_cell(cell, "F4F6F9")
        _set_cell_margins(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_after = 0
        run = paragraph.add_run(label)
        _set_run_font(run, bold=True, color="0B2545", size=10)
    for row_values in rows:
        row = table.add_row()
        row_properties = row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        cant_split.set(qn("w:val"), "true")
        row_properties.append(cant_split)
        cells = row.cells
        for cell, value in zip(cells, row_values, strict=True):
            _set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = 0
            run = paragraph.add_run(value)
            _set_run_font(run, size=10)
    _set_table_geometry(table, widths)
    return table


def _add_bullet(document: Any, text: str) -> None:
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.add_run(text)


def _add_compact_source_bullet(document: Any, text: str) -> None:
    """Add a readable, space-efficient source citation for the packet appendix."""
    from docx.shared import Pt

    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run(text)
    _set_run_font(run, size=9)


def _add_numbered(document: Any, text: str) -> None:
    paragraph = document.add_paragraph(style="List Number")
    paragraph.add_run(text)


def _render_docx(packet: Mapping[str, Any], path: Path) -> None:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    document = Document()
    _configure_document_styles(document)
    document.core_properties.title = str(packet["opportunity"]["opportunity_name"])
    document.core_properties.subject = "Grant application preparation and human-review packet"
    document.core_properties.author = "BrokenGrowthMinistries - GrantBot Pro"

    section = document.sections[0]
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_run = header.add_run("BrokenGrowthMinistries | GrantBot Human-Review Draft")
    _set_run_font(header_run, size=8.5, color="666666")
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run("Not authorized for submission, certification, signature, payment, or financial commitment")
    _set_run_font(footer_run, size=8, color="777777", italic=True)

    applicant = document.add_paragraph()
    applicant.alignment = WD_ALIGN_PARAGRAPH.CENTER
    applicant.paragraph_format.space_after = Pt(8)
    applicant_run = applicant.add_run(str(packet["organization_name"]))
    _set_run_font(applicant_run, size=12, color="555555", bold=True)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(4)
    title_run = title.add_run(str(packet["opportunity"]["opportunity_name"]))
    _set_run_font(title_run, size=24, color="000000", bold=True)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(18)
    subtitle_run = subtitle.add_run("Application Preparation and Human-Review Packet")
    _set_run_font(subtitle_run, size=14, color="555555")

    request = packet["opportunity"].get("request") or {}
    metadata_rows = [
        ("Funder", str(packet["opportunity"]["funder"])),
        ("Funding type", str(packet["opportunity"]["funding_type"])),
        ("Deadline", str(packet["opportunity"].get("deadline") or "Not stated")),
        ("Requested amount", str(request.get("amount") or "Not yet approved")),
        ("Readiness", f"{packet['readiness_score']}% - {_display_status(packet['status'])}"),
        ("Official source", str(packet["opportunity"]["source_url"])),
    ]
    _add_table(document, ("Application item", "Prepared information"), metadata_rows, (2700, 6660))

    document.add_paragraph()
    notice = _add_table(
        document,
        ("Control notice",),
        (("The live portal questions must be checked after login. GrantBot has not submitted this request and cannot make certifications or financial commitments.",),),
        (9360,),
    )
    _shade_cell(notice.rows[1].cells[0], "FFF8E8")

    document.add_page_break()
    document.add_heading("Eligibility Review", level=1)
    eligibility_rows = []
    for item in packet["eligibility"]:
        evidence = "; ".join(
            f"{_display_key(entry['fact_key'])}: {_value_text(entry['value'])}"
            for entry in item["evidence"]
        ) or item["action"]
        eligibility_rows.append((str(item["label"]), _display_status(str(item["status"])), evidence))
    _add_table(document, ("Requirement", "Status", "Evidence or action"), eligibility_rows, (2600, 1500, 5260))

    if packet["opportunity"]["external_evidence"]:
        document.add_heading("Verified External Need Evidence", level=1)
        evidence_rows = [
            (
                str(item["label"]),
                str(item.get("source_values") or ""),
                str(item.get("caveat") or ""),
            )
            for item in packet["opportunity"]["external_evidence"]
        ]
        _add_table(
            document,
            ("Evidence", "Source values", "Scope caveat"),
            evidence_rows,
            (2100, 4300, 2960),
        )

    document.add_heading("Registration and Portal Fields", level=1)
    portal_groups = [
        (
            "Tax and legal identity",
            {"registered_ein_answer", "ein", "legal_name", "organization_type", "mission_statement"},
        ),
        (
            "Physical and mailing addresses",
            {
                "physical_address", "physical_city", "physical_state", "physical_county", "physical_zip",
                "mailing_address", "mailing_city", "mailing_state", "mailing_zip",
            },
        ),
        (
            "Organization channels and social media",
            {"organization_phone", "organization_email", "website", "social_media_status", "social_media_urls"},
        ),
        (
            "Authorized applicant contact",
            {"contact_first_name", "contact_last_name", "contact_title", "contact_phone", "contact_email"},
        ),
    ]
    for group_label, group_keys in portal_groups:
        group_items = [
            item for item in packet["portal_fields"] if item["key"] in group_keys
        ]
        if not group_items:
            continue
        document.add_heading(group_label, level=2)
        portal_rows = [
            (
                str(item["label"]),
                _display_status(str(item["status"])),
                _value_text(item["value"]) or "Not supplied - action required",
                str(item["instructions"] or "Verify against the live portal."),
            )
            for item in group_items
        ]
        _add_table(
            document,
            ("Field", "Status", "Prepared value", "Entry instruction"),
            portal_rows,
            (2100, 1200, 3100, 2960),
        )

    document.add_heading("Prepared Narratives", level=1)
    for item in packet["sections"]:
        document.add_heading(str(item["label"]), level=2)
        detail = document.add_paragraph()
        detail.paragraph_format.space_after = Pt(5)
        detail_run = detail.add_run(f"Status: {_display_status(item['status'])} | Words: {item['word_count']}")
        _set_run_font(detail_run, size=9, color="666666", bold=True)
        prompt = document.add_paragraph()
        prompt_run = prompt.add_run("Preparation prompt: ")
        _set_run_font(prompt_run, bold=True, size=10)
        prompt_value = prompt.add_run(str(item["prompt"]))
        _set_run_font(prompt_value, italic=True, size=10, color="444444")
        narrative = document.add_paragraph(str(item.get("draft") or "No narrative generated because required information is missing."))
        narrative.paragraph_format.keep_together = False
        if item["missing_fact_keys"]:
            missing = document.add_paragraph()
            missing_run = missing.add_run(
                "Information still required: "
                + ", ".join(_display_key(key) for key in item["missing_fact_keys"])
            )
            _set_run_font(missing_run, color="9B1C1C", bold=True, size=10)

    document.add_heading("Human-Required Gates", level=1)
    if packet["hard_gates"]:
        for item in packet["hard_gates"]:
            _add_bullet(document, str(item))
    else:
        document.add_paragraph("No unresolved hard gates were recorded. Authorized final review is still required.")

    document.add_heading("Document Preparation Checklist", level=1)
    for item in packet["attachments"]:
        _add_bullet(document, f"{item['label']} - {_display_status(item['status'])}")

    document.add_heading("Submission Sequence", level=1)
    for step in packet["submission_steps"]:
        _add_numbered(document, str(step))

    document.add_heading("Official Sources", level=1)
    for source in packet["opportunity"]["official_sources"]:
        _add_compact_source_bullet(
            document,
            f"{source.get('title') or source['url']}: {source['url']} - {_value_text(source.get('supports'))}",
        )

    path.parent.mkdir(parents=True, exist_ok=True)
    fd, temporary = tempfile.mkstemp(
        prefix=path.stem + "-",
        suffix=".docx.tmp",
        dir=str(path.parent),
    )
    os.close(fd)
    try:
        document.save(temporary)
        os.replace(temporary, path)
    finally:
        if os.path.exists(temporary):
            os.unlink(temporary)


def build_production_application_packet(
    opportunity: Mapping[str, Any],
    *,
    output_root: Path | None = None,
    draft_mode: str = "template",
    supplied_answers: Mapping[str, str] | None = None,
    include_docx: bool = True,
    run_adversarial_review: bool = False,
) -> dict[str, Any]:
    normalized = _validate_opportunity(opportunity)
    facts = canonical_application_facts()
    fields = _resolve_portal_fields(normalized["required_portal_fields"], facts)
    eligibility = _eligibility_results(normalized["eligibility_checks"], facts)
    sections = _draft_sections(
        normalized,
        facts,
        draft_mode=draft_mode,
        supplied_answers=supplied_answers,
        run_adversarial_review=run_adversarial_review,
    )
    attachments = _attachment_results(normalized["attachment_checklist"], facts)
    deadline_review = _deadline_review(normalized.get("deadline"))

    hard_gates: list[str] = []
    portal_gate_groups: dict[str, list[ResolvedField]] = {}

    def portal_gate_group(field: ResolvedField) -> str:
        if field.key in {"registered_ein_answer", "ein"}:
            return "Tax identity"
        if field.key.startswith("physical_"):
            return "Organization physical address"
        if field.key.startswith("mailing_"):
            return "Organization mailing address"
        if field.key.startswith("contact_"):
            return "Authorized applicant contact"
        if field.key.startswith("social_media"):
            return "Social media disclosure"
        if field.key in {"organization_phone", "organization_email", "website"}:
            return "Organization contact channels"
        return "Organization registration"

    for field in fields:
        if field.hard_gate and field.status != "READY":
            portal_gate_groups.setdefault(portal_gate_group(field), []).append(field)
    for group, items in portal_gate_groups.items():
        labels = ", ".join(item.label for item in items)
        instructions = _dedupe([item.instructions for item in items])
        action = " ".join(instructions)
        hard_gates.append(
            f"Portal - {group}: complete {labels}. {action or 'Verify against authoritative records.'}"
        )
    for item in eligibility:
        if item.status != "PASS":
            hard_gates.append(f"Eligibility - {item.label}: {item.action}")
    for section in sections:
        if section.status == "NEEDS_INFORMATION":
            hard_gates.append(
                f"Narrative - {section.label}: supply "
                f"{', '.join(_display_key(key) for key in section.missing_fact_keys) or 'the missing source facts'}"
            )
    for item in attachments:
        if item["required"] and item["status"] != "AVAILABLE":
            hard_gates.append(f"Document - {item['label']}: {item['instructions'] or item['status']}")
    request = normalized.get("request") or {}
    if request.get("amount") in (None, ""):
        hard_gates.append("Approve the exact request amount and a line-item use-of-funds budget.")
    if deadline_review["status"] == "EXPIRED":
        hard_gates.append(f"The published deadline passed on {deadline_review['deadline']}.")
    hard_gates = _dedupe(hard_gates)

    readiness = _readiness_score(fields, eligibility, sections, attachments)
    eligibility_failure = any(item.status == "FAIL" for item in eligibility)
    if eligibility_failure or deadline_review["status"] == "EXPIRED":
        status = "HOLD"
    elif hard_gates:
        status = "NEEDS_USER"
    elif all(item.status == "READY_FOR_HUMAN_REVIEW" for item in sections):
        status = "READY_FOR_REVIEW"
    else:
        status = "REVISION_REQUIRED"

    organization_name = _fact_value(_fact_index(facts), "legal_name", "organization_name")
    created_at = datetime.now(timezone.utc).isoformat()
    payload: dict[str, Any] = {
        "schema_version": "30.1",
        "packet_id": normalized["opportunity_id"],
        "created_at": created_at,
        "organization_name": organization_name,
        "opportunity": normalized,
        "status": status,
        "readiness_score": readiness,
        "deadline_review": deadline_review,
        "eligibility": [item.to_dict() for item in eligibility],
        "portal_fields": [item.to_dict() for item in fields],
        "sections": [item.to_dict() for item in sections],
        "attachments": attachments,
        "hard_gates": hard_gates,
        "submission_steps": list(normalized["submission_steps"]),
        "fact_snapshot_sha256": hashlib.sha256(_json_text(facts).encode("utf-8")).hexdigest(),
        "draft_mode": draft_mode,
        "human_review_required": True,
        "submission_performed": False,
        "external_submission_enabled": False,
        "safe_to_submit": False,
    }

    root = Path(output_root).expanduser().resolve() if output_root else application_root()
    folder = root / normalized["opportunity_id"]
    folder.mkdir(parents=True, exist_ok=True)
    json_path = folder / "application_packet.json"
    markdown_path = folder / "APPLICATION_DRAFT.md"
    checklist_path = folder / "HUMAN_REVIEW_CHECKLIST.md"
    source_path = folder / "source_manifest.json"
    docx_path = folder / "APPLICATION_PACKET.docx"
    payload["artifacts"] = {
        "folder": str(folder),
        "json": str(json_path),
        "markdown": str(markdown_path),
        "checklist": str(checklist_path),
        "source_manifest": str(source_path),
        "docx": str(docx_path) if include_docx else None,
    }
    _atomic_write(json_path, json.dumps(payload, indent=2, ensure_ascii=False, sort_keys=True) + "\n")
    _atomic_write(markdown_path, _render_application_markdown(payload))
    _atomic_write(checklist_path, _render_checklist_markdown(payload))
    _atomic_write(
        source_path,
        json.dumps(
            {
                "created_at": created_at,
                "official_sources": normalized["official_sources"],
                "external_evidence": normalized["external_evidence"],
                "fact_snapshot_sha256": payload["fact_snapshot_sha256"],
                "control": "Portal prompts and requirements must be verified after login.",
            },
            indent=2,
            ensure_ascii=False,
            sort_keys=True,
        ) + "\n",
    )
    if include_docx:
        _render_docx(payload, docx_path)
    return payload
