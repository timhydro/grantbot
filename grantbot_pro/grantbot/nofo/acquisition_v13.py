from __future__ import annotations

import html
import json
import re
import ssl
import urllib.parse
import urllib.request
from dataclasses import asdict, dataclass
from html.parser import HTMLParser
from io import BytesIO
from pathlib import Path
from typing import Any

from docx import Document as DocxDocument
from pypdf import PdfReader

from grantbot.discovery.grants_gov import fetch_opportunity
from grantbot.evidence.document_intelligence import AUTHORITY_RANKS, classify_authority
from grantbot.evidence.repository import ingest_document, list_requirements
from grantbot.nofo.full_detail import get_full_nofo_intelligence


MAX_DOWNLOAD_BYTES = 30 * 1024 * 1024
TIMEOUT_SECONDS = 30
MAX_DOCUMENTS = 40
MAX_BLUEPRINT_TEXT_BYTES = 2_000_000
ALLOWED_HOST_SUFFIXES = (".gov", "grants.gov", "simpler.grants.gov")

QUESTION_RE = re.compile(
    r"\b(describe|explain|provide|discuss|identify|demonstrate|summarize|"
    r"detail|outline|justify|address|state|specify|document)\b",
    re.I,
)


@dataclass(frozen=True, slots=True)
class ApplicationBlueprint:
    opportunity_id: str
    opportunity_number: str
    title: str
    funder: str
    acquisition_status: str
    source_urls: list[str]
    application_questions: list[str]
    requirements: list[str]
    scoring_criteria: list[str]
    budget_requirements: list[str]
    match_requirements: list[str]
    submission_requirements: list[str]
    required_attachments: list[str]
    eligibility: list[str]
    award_ceiling: float | None
    award_floor: float | None
    cost_sharing: bool | None
    warnings: list[str]
    output_path: str

    def to_dict(self) -> dict[str, Any]:
        return asdict(self)


@dataclass(frozen=True, slots=True)
class AcquiredDocument:
    url: str
    content_type: str
    title: str
    authority_type: str
    authority_rank: int
    pages: list[str]


class _HTMLTextParser(HTMLParser):
    def __init__(self) -> None:
        super().__init__()
        self.parts: list[str] = []
        self.links: list[str] = []
        self._ignore = 0

    def handle_starttag(
        self,
        tag: str,
        attrs: list[tuple[str, str | None]],
    ) -> None:
        low = tag.lower()
        if low in {"script", "style", "noscript", "svg"}:
            self._ignore += 1
        if low == "a":
            for key, value in attrs:
                if key.lower() == "href" and value:
                    self.links.append(value)

    def handle_endtag(self, tag: str) -> None:
        if tag.lower() in {"script", "style", "noscript", "svg"} and self._ignore:
            self._ignore -= 1

    def handle_data(self, data: str) -> None:
        if self._ignore:
            return
        value = " ".join(data.split()).strip()
        if value:
            self.parts.append(value)


def _safe_url(url: str) -> bool:
    try:
        parsed = urllib.parse.urlparse(url)
    except ValueError:
        return False
    if parsed.scheme != "https":
        return False
    host = (parsed.hostname or "").lower()
    return any(
        host == suffix.lstrip(".") or host.endswith(suffix)
        for suffix in ALLOWED_HOST_SUFFIXES
    )


def _dedupe(values: list[str], limit: int = 300) -> list[str]:
    result: list[str] = []
    seen: set[str] = set()
    for raw in values:
        value = " ".join(str(raw).split()).strip()
        key = value.casefold()
        if not value or key in seen:
            continue
        seen.add(key)
        result.append(value)
        if len(result) >= limit:
            break
    return result


def _walk_strings(value: Any) -> list[str]:
    result: list[str] = []
    if isinstance(value, str):
        result.append(value)
    elif isinstance(value, dict):
        for nested in value.values():
            result.extend(_walk_strings(nested))
    elif isinstance(value, list):
        for nested in value:
            result.extend(_walk_strings(nested))
    return result


def _request(url: str) -> tuple[bytes, str, str]:
    request = urllib.request.Request(
        url,
        headers={
            "User-Agent": "GrantBotPro/20.0",
            "Accept": (
                "application/pdf,application/vnd.openxmlformats-officedocument."
                "wordprocessingml.document,text/html,text/plain,*/*;q=0.2"
            ),
        },
    )
    with urllib.request.urlopen(
        request,
        timeout=TIMEOUT_SECONDS,
        context=ssl.create_default_context(),
    ) as response:
        final_url = response.geturl()
        if not _safe_url(final_url):
            raise ValueError(f"Unsafe redirect: {final_url}")
        content_type = response.headers.get_content_type() or "application/octet-stream"
        data = response.read(MAX_DOWNLOAD_BYTES + 1)
        if len(data) > MAX_DOWNLOAD_BYTES:
            raise ValueError("NOFO document exceeds download size limit")
        return data, content_type, final_url


def _extract_document(
    data: bytes,
    content_type: str,
    url: str,
) -> tuple[list[str], list[str]]:
    path = urllib.parse.urlparse(url).path.lower()
    if content_type == "application/pdf" or path.endswith(".pdf"):
        reader = PdfReader(BytesIO(data))
        pages = [(page.extract_text() or "").strip() for page in reader.pages]
        return pages, []

    if (
        content_type
        == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        or path.endswith(".docx")
    ):
        document = DocxDocument(BytesIO(data))
        text = "\n".join(
            paragraph.text.strip()
            for paragraph in document.paragraphs
            if paragraph.text.strip()
        )
        return [text], []

    if content_type == "text/html" or path.endswith((".html", ".htm")):
        parser = _HTMLTextParser()
        parser.feed(data.decode("utf-8", errors="replace"))
        links = [urllib.parse.urljoin(url, href) for href in parser.links]
        return ["\n".join(parser.parts)], links

    if content_type.startswith("text/") or path.endswith(".txt"):
        return [data.decode("utf-8", errors="replace")], []

    return [], []


def _document_title(url: str, fallback: str) -> str:
    path = urllib.parse.unquote(urllib.parse.urlparse(url).path)
    filename = Path(path).name.strip()
    return filename or fallback


def _resolve_authority(
    *,
    url: str,
    title: str,
    content_type: str,
) -> str:
    resolved = classify_authority(
        source_url=url,
        title=title,
        document_type=content_type,
    )
    if resolved == "INFORMATIONAL_WEBPAGE" and (
        content_type == "application/pdf" or url.lower().endswith(".pdf")
    ):
        filename = Path(urllib.parse.urlparse(url).path).name.lower()
        if any(token in filename for token in ("notice", "nofo", "foa", "funding-opportunity")):
            return "PRIMARY_NOFO"
        if "instruction" in filename or "application-guide" in filename:
            return "APPLICATION_INSTRUCTIONS"
    return resolved


def _lines(text: str) -> list[str]:
    result: list[str] = []
    for raw in text.replace("\r", "\n").split("\n"):
        value = " ".join(raw.split()).strip()
        if 10 <= len(value) <= 1200:
            result.append(value)
    return result


def _validated_questions(lines: list[str]) -> list[str]:
    noise = (
        "questions or comments",
        "can you talk about",
        "did you read",
        "do you know",
        "how did you",
        "who else",
        "office hours",
        "submit questions",
        "comments or questions",
    )
    result: list[str] = []
    for raw in lines:
        text = " ".join(str(raw).split()).strip()
        low = text.lower()
        if len(text) < 8 or len(text) > 700:
            continue
        if any(item in low for item in noise):
            continue
        starts_as_prompt = any(
            low.startswith(prefix)
            for prefix in (
                "describe ",
                "explain ",
                "provide ",
                "identify ",
                "demonstrate ",
                "discuss ",
                "summarize ",
                "outline ",
                "justify ",
                "address ",
            )
        )
        if ("?" in text or starts_as_prompt) and QUESTION_RE.search(text):
            result.append(text)
    return _dedupe(result, limit=250)


def _authoritative_documents(documents: list[AcquiredDocument]) -> list[AcquiredDocument]:
    strong = [item for item in documents if item.authority_rank <= 6]
    return strong or documents


def _requirement_texts(
    requirements: list[dict[str, Any]],
    *categories: str,
    limit: int = 250,
) -> list[str]:
    allowed = {category.upper() for category in categories}
    values = [
        str(item.get("text", ""))
        for item in requirements
        if not allowed or str(item.get("category", "")).upper() in allowed
    ]
    return _dedupe(values, limit=limit)


def _queue_link(link: str) -> bool:
    low = link.lower()
    return _safe_url(link) and any(
        token in low
        for token in (
            ".pdf",
            ".docx",
            "nofo",
            "notice",
            "attachment",
            "download",
            "instructions",
            "application-guide",
            "amendment",
            "faq",
        )
    )


def acquire_blueprint(
    opportunity_id: str,
    *,
    manual_urls: list[str] | None = None,
) -> ApplicationBlueprint:
    opportunity_id = opportunity_id.strip()
    if not opportunity_id or len(opportunity_id) > 80:
        raise ValueError("Invalid opportunity_id")

    raw = fetch_opportunity(opportunity_id)
    if not isinstance(raw, dict):
        raise RuntimeError("Invalid Grants.gov opportunity payload")
    full = get_full_nofo_intelligence(opportunity_id)

    urls = [
        f"https://www.grants.gov/search-results-detail/"
        f"{urllib.parse.quote(opportunity_id, safe='')}"
    ]
    for value in _walk_strings(raw):
        value = html.unescape(value.strip())
        if value.startswith("https://") and _safe_url(value):
            urls.append(value)
    for value in manual_urls or []:
        value = value.strip()
        if not _safe_url(value):
            raise ValueError(
                f"Manual NOFO URL must be an HTTPS government source: {value}"
            )
        urls.append(value)

    queue = _dedupe(urls, limit=60)
    visited: set[str] = set()
    acquired: list[AcquiredDocument] = []
    warnings: list[str] = []

    while queue and len(visited) < MAX_DOCUMENTS:
        url = queue.pop(0)
        if url in visited or not _safe_url(url):
            continue
        visited.add(url)
        try:
            data, content_type, final_url = _request(url)
            pages, links = _extract_document(data, content_type, final_url)
            pages = [page for page in pages if page.strip()]
            if pages:
                title = _document_title(final_url, full.title)
                authority = _resolve_authority(
                    url=final_url,
                    title=title,
                    content_type=content_type,
                )
                document = AcquiredDocument(
                    url=final_url,
                    content_type=content_type,
                    title=title,
                    authority_type=authority,
                    authority_rank=AUTHORITY_RANKS[authority],
                    pages=pages,
                )
                acquired.append(document)
                try:
                    ingest_document(
                        opportunity_id=opportunity_id,
                        title=title,
                        pages=pages,
                        source_url=final_url,
                        document_type=content_type,
                        authority_type=authority,
                    )
                except Exception as exc:
                    warnings.append(
                        f"Evidence indexing warning for {final_url}: "
                        f"{type(exc).__name__}: {exc}"
                    )
            for link in links:
                if _queue_link(link) and link not in visited and link not in queue:
                    queue.append(link)
        except Exception as exc:
            warnings.append(f"{url}: {type(exc).__name__}: {exc}")

    authoritative_docs = _authoritative_documents(acquired)
    authoritative_text = "\n\n".join(
        page
        for document in authoritative_docs
        for page in document.pages
    )[:MAX_BLUEPRINT_TEXT_BYTES]
    questions = _validated_questions(_lines(authoritative_text))

    try:
        evidence_requirements = list_requirements(
            opportunity_id=opportunity_id,
            authoritative_only=True,
        )
    except Exception as exc:
        evidence_requirements = []
        warnings.append(
            f"Requirement evidence retrieval warning: {type(exc).__name__}: {exc}"
        )

    requirements = _requirement_texts(evidence_requirements)
    scoring_criteria = _requirement_texts(evidence_requirements, "SCORING")
    budget_requirements = _requirement_texts(evidence_requirements, "BUDGET")
    match_requirements = _requirement_texts(evidence_requirements, "MATCH")
    submission_requirements = _requirement_texts(evidence_requirements, "SUBMISSION")
    required_attachments = _requirement_texts(evidence_requirements, "ATTACHMENT")

    quality_count = sum(
        bool(value)
        for value in (
            questions,
            requirements,
            scoring_criteria,
            budget_requirements,
            submission_requirements,
            required_attachments,
            full.eligibility,
        )
    )
    acquired_text = any(document.pages for document in acquired)
    strong_sources = [item for item in acquired if item.authority_rank <= 6]
    if not acquired_text:
        acquisition_status = "METADATA_ONLY"
    elif strong_sources and quality_count >= 5 and requirements and full.eligibility:
        acquisition_status = "FULL_NOFO_VALIDATED"
    else:
        acquisition_status = "FULL_TEXT_ACQUIRED"

    if not questions:
        warnings.append(
            "No application questions were extracted from authoritative source text."
        )
    if acquired and not strong_sources:
        warnings.append(
            "No primary NOFO, amendment, application instructions, form, FAQ, or official guidance "
            "was confidently identified; requirement extraction used the best available government text."
        )

    root = (
        Path(__file__).resolve().parents[2]
        / "data"
        / "nofo_blueprints"
        / re.sub(r"[^A-Za-z0-9._-]+", "_", opportunity_id)
    )
    root.mkdir(parents=True, exist_ok=True)
    output_path = root / "application_blueprint.json"

    blueprint = ApplicationBlueprint(
        opportunity_id=opportunity_id,
        opportunity_number=full.opportunity_number,
        title=full.title,
        funder=full.funder,
        acquisition_status=acquisition_status,
        source_urls=_dedupe([document.url for document in acquired]),
        application_questions=questions,
        requirements=requirements,
        scoring_criteria=scoring_criteria,
        budget_requirements=budget_requirements,
        match_requirements=match_requirements,
        submission_requirements=submission_requirements,
        required_attachments=required_attachments,
        eligibility=full.eligibility,
        award_ceiling=full.award_ceiling,
        award_floor=full.award_floor,
        cost_sharing=full.cost_sharing,
        warnings=_dedupe(warnings, limit=100),
        output_path=str(output_path),
    )
    output_path.write_text(
        json.dumps(blueprint.to_dict(), indent=2, ensure_ascii=False) + "\n",
        encoding="utf-8",
    )
    return blueprint


def load_blueprint(opportunity_id: str) -> dict[str, Any]:
    safe_id = re.sub(
        r"[^A-Za-z0-9._-]+",
        "_",
        opportunity_id.strip(),
    )
    path = (
        Path(__file__).resolve().parents[2]
        / "data"
        / "nofo_blueprints"
        / safe_id
        / "application_blueprint.json"
    )
    if not path.exists():
        raise FileNotFoundError(str(path))
    data = json.loads(path.read_text(encoding="utf-8"))
    if not isinstance(data, dict):
        raise RuntimeError("Stored NOFO blueprint is invalid")
    return data
